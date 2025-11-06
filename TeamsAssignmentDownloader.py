import os
from typing import List, Optional, Tuple
import shutil

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


"""
Local Teams assignment file collector:
  <Root>\\<Student Name>\\<Assignment Name>\\<files>
Copies files to output, renaming to student's name and preserving extensions.
"""


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def sanitize_filename(name: str) -> str:
    invalid = '<>:"/\\|?*'  # windows invalid characters
    sanitized = "".join(" " if c in invalid else c for c in name)
    # Replace underscores with spaces
    sanitized = sanitized.replace("_", " ")
    # Collapse multiple spaces into one
    while "  " in sanitized:
        sanitized = sanitized.replace("  ", " ")
    sanitized = sanitized.strip().rstrip(".")
    return sanitized or "unnamed"


def choose_unique_filename(base_name: str, ext: str, folder: str) -> str:
    candidate = f"{base_name}{ext}"
    path = os.path.join(folder, candidate)
    if not os.path.exists(path):
        return path
    i = 2
    while True:
        candidate = f"{base_name} ({i}){ext}"
        path = os.path.join(folder, candidate)
        if not os.path.exists(path):
            return path
        i += 1


def add_filename_to_docx_header(docx_path: str, filename: str) -> bool:
    """Add the filename to the header of a DOCX document.
    
    Args:
        docx_path: Path to the DOCX file
        filename: Filename (without extension) to add to header
        
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        return False
    
    try:
        doc = Document(docx_path)
        
        # Access the first section's header (default header)
        section = doc.sections[0]
        header = section.header
        
        # Get or create the first paragraph in the header
        if len(header.paragraphs) == 0:
            paragraph = header.add_paragraph()
        else:
            paragraph = header.paragraphs[0]
        
        # Clear existing content and add filename
        paragraph.clear()
        run = paragraph.add_run(filename)
        run.font.size = Pt(10)
        
        # Save the document
        doc.save(docx_path)
        return True
    except Exception:
        # If anything goes wrong, return False (file will still be copied)
        return False


def copy_from_local_folder(
    turned_in_path: str, 
    output: str, 
    assignment_subfolder: Optional[str],
    prefix: Optional[str] = None,
    suffix: Optional[str] = None,
    add_to_header: bool = False
) -> Tuple[int, int]:
    """Copy files from a local synced 'Turned in' folder.

    The expected structure is:
      <Turned in>/
        <Student Name>/
          file1.ext
          file2.ext

    Args:
        turned_in_path: Path to the root folder containing student submissions
        output: Output folder path
        assignment_subfolder: Optional assignment subfolder name to filter by
        prefix: Optional prefix to add before student name in filename
        suffix: Optional suffix to add after student name in filename
        add_to_header: If True, add filename to DOCX document headers

    Returns (total_files_found, total_files_copied)
    """
    if not os.path.isdir(turned_in_path):
        raise RuntimeError(f"Path does not exist or is not a directory: {turned_in_path}")

    total_found = 0
    total_copied = 0
    for entry in os.scandir(turned_in_path):
        if not entry.is_dir():
            continue
        student_name = sanitize_filename(entry.name)
        student_dir = entry.path

        # Build filename with prefix and suffix
        safe_prefix = sanitize_filename(prefix).strip() if prefix and prefix.strip() else ""
        safe_suffix = sanitize_filename(suffix).strip() if suffix and suffix.strip() else ""
        
        # Construct base name: prefix + student_name + suffix (using spaces, not underscores)
        if safe_prefix and safe_suffix:
            base_name = f"{safe_prefix} {student_name} {safe_suffix}"
        elif safe_prefix:
            base_name = f"{safe_prefix} {student_name}"
        elif safe_suffix:
            base_name = f"{student_name} {safe_suffix}"
        else:
            base_name = student_name

        # If an assignment subfolder was specified, prefer that
        roots_to_walk: List[str] = []
        if assignment_subfolder:
            maybe = os.path.join(student_dir, assignment_subfolder)
            if os.path.isdir(maybe):
                roots_to_walk.append(maybe)
            else:
                # If the specified assignment folder doesn't exist for this student, skip
                continue
        else:
            roots_to_walk.append(student_dir)

        for root in roots_to_walk:
            for _dirpath, _dirs, files in os.walk(root):
                for fname in files:
                    total_found += 1
                    _, ext = os.path.splitext(fname)
                    dest_path = choose_unique_filename(base_name, ext, output)
                    src_path = os.path.join(_dirpath, fname)
                    shutil.copy2(src_path, dest_path)
                    
                    # Add filename to DOCX header if requested
                    if add_to_header and ext.lower() == '.docx':
                        add_filename_to_docx_header(dest_path, base_name)
                    
                    total_copied += 1
    return total_found, total_copied


def collect_assignment_subfolders(turned_in_path: str) -> List[str]:
    """Scan student folders one level deep and collect unique assignment subfolder names."""
    names: set[str] = set()
    if not os.path.isdir(turned_in_path):
        return []
    for entry in os.scandir(turned_in_path):
        if not entry.is_dir():
            continue
        student_dir = entry.path
        try:
            for sub in os.scandir(student_dir):
                if sub.is_dir():
                    names.add(sub.name)
        except PermissionError:
            continue
    return sorted(names)


def main() -> None:
    print("Microsoft Teams Assignment Downloader")
    print("-----------------------------------")

    print("Provide the path to the submitted files root. Structure:")
    print("  <Root>\\<Student Name>\\<Assignment Name>\\<files>")
    turned_in_path = ""
    while not turned_in_path:
        p = input("Enter path to submitted files folder: ").strip()
        if not p:
            print("A path is required.")
            continue
        if not os.path.isdir(p):
            print("That path does not exist or is not a directory. Try again.")
            continue
        turned_in_path = p

    # Detect assignment subfolders
    assignments = collect_assignment_subfolders(turned_in_path)
    selected_assignment: Optional[str] = None
    if assignments:
        print("Detected assignment folders:")
        for idx, name in enumerate(assignments, start=1):
            print(f"  {idx}. {name}")
        while True:
            sel = input("Select assignment number to extract (or type exact name): ").strip()
            if sel.isdigit():
                i = int(sel)
                if 1 <= i <= len(assignments):
                    selected_assignment = assignments[i - 1]
                    break
            elif sel:
                if sel in assignments:
                    selected_assignment = sel
                    break
                else:
                    print("Name not found in detected list. Please choose a number or exact name.")
                    continue
            else:
                print("A selection is required.")
    else:
        print("No assignment subfolders detected; will copy files directly under each student folder.")
        selected_assignment = None

    # Create default output folder based on assignment name
    program_root = os.path.dirname(os.path.abspath(__file__))
    if selected_assignment:
        safe_assignment_name = sanitize_filename(selected_assignment)
        default_out = os.path.join(program_root, "downloads", safe_assignment_name)
    else:
        default_out = os.path.join(program_root, "downloads", "unnamed_assignment")
    
    output = input(f"Enter output folder path [{default_out}]: ").strip() or default_out
    ensure_dir(output)

    # Ask for prefix and suffix
    prefix = input("Enter prefix to add before student name (optional, press Enter to skip): ").strip() or None
    suffix = input("Enter suffix to add after student name (optional, press Enter to skip): ").strip() or None

    # Ask if user wants to add filename to DOCX headers
    add_to_header = False
    if DOCX_AVAILABLE:
        header_input = input("Add filename to DOCX document headers? [y/N]: ").strip().lower()
        add_to_header = header_input in ("y", "yes")
    else:
        header_input = input("Add filename to DOCX document headers? [y/N] (requires python-docx): ").strip().lower()
        if header_input in ("y", "yes"):
            print("Warning: python-docx is not installed. Headers will not be modified.")
            print("Install it with: pip install python-docx")

    total_found, total_copied = copy_from_local_folder(turned_in_path, output, selected_assignment, prefix, suffix, add_to_header)
    print(f"Done. Files found: {total_found}, copied: {total_copied}.")
    return


if __name__ == "__main__":
    main()


